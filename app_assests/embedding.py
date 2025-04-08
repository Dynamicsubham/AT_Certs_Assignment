import fitz
from docx import Document as DocxDocument
import pandas as pd
import os
from io import BytesIO

from sentence_transformers import SentenceTransformer

model = SentenceTransformer('all-MiniLM-L6-v2')


from io import BytesIO
from typing import Union

def extract_text(file_content: Union[bytes, BytesIO], filename: str) -> str:
    try:
        if isinstance(file_content, bytes):
            file_content = BytesIO(file_content)

        if filename.endswith(".pdf"):
            from PyPDF2 import PdfReader
            reader = PdfReader(file_content)
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        elif filename.endswith(".docx"):
            import docx
            doc = docx.Document(file_content)
            return "\n".join(p.text for p in doc.paragraphs)

        elif filename.endswith(".xlsx"):
            from openpyxl import load_workbook
            wb = load_workbook(file_content, data_only=True)
            text = ""
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    text += " ".join(str(cell) if cell is not None else "" for cell in row) + "\n"
            return text

        elif filename.endswith(".xls"):
            import xlrd
            wb = xlrd.open_workbook(file_contents=file_content.read())
            text = ""
            for sheet in wb.sheets():
                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    text += " ".join(str(cell) for cell in row) + "\n"
            return text

        else:
            raise ValueError("Unsupported file format")

    except Exception as e:
        raise RuntimeError(f"Failed to extract text from {filename}: {str(e)}")

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def extract_text_from_excel(file_path: str) -> str:
    dfs = pd.read_excel(file_path, sheet_name=None)
    combined_text = ""
    for name, df in dfs.items():
        combined_text += f"\nSheet: {name}\n"
        combined_text += df.astype(str).to_string(index=False)
    return combined_text

def get_embedding(text: str):
    return model.encode(text).tolist()
